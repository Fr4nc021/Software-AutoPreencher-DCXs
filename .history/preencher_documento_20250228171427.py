import tkinter as tk
from tkinter import messagebox
from docx import Document

def gerar_documentos(dados):
    doc = Document()
    doc.add_heading('Dados do cliente', 0)
    for key, value in dados.items():
        doc.add_paragraph(f"{key}: {value}")

        doc.save(f"documento_{dados['Nome'].replace(' ', '_')}_{dados['RG']}.docx")
        messagebox.showinfo("Sucesso", "Documento gerado com sucesso!")

# Função para coletar e exibir os dados preenchidos
def coletar_dados():
    dados = {
        "Nome": entry_nome.get(),
        "RG": entry_rg.get(),
        "CPF": entry_cpf.get(),
        "Endereço Completo": entry_endereco.get(),
        "Cidade de Nascimento": entry_cidade.get(),
        "Estado de Nascimento": entry_estado.get(),
        "Data de Nascimento": entry_data.get(),
        "Nome do Pai": entry_pai.get(),
        "Nome da Mãe": entry_mae.get(),
        "Data de Expedição do RG": entry_data_exp_rg.get(),
    }
    
    # Exibir os dados preenchidos em uma mensagem
    dados_str = "\n".join([f"{key}: {value}" for key, value in dados.items()])
    messagebox.showinfo("Dados Preenchidos", dados_str)

# Criando a interface gráfica com Tkinter
root = tk.Tk()
root.title("Preenchimento de Dados")

# Definindo os campos de entrada
label_nome = tk.Label(root, text="Nome:")
label_nome.grid(row=0, column=0, sticky="w", padx=10, pady=5)
entry_nome = tk.Entry(root)
entry_nome.grid(row=0, column=1, padx=10, pady=5)

label_rg = tk.Label(root, text="RG:")
label_rg.grid(row=1, column=0, sticky="w", padx=10, pady=5)
entry_rg = tk.Entry(root)
entry_rg.grid(row=1, column=1, padx=10, pady=5)

label_cpf = tk.Label(root, text="CPF:")
label_cpf.grid(row=2, column=0, sticky="w", padx=10, pady=5)
entry_cpf = tk.Entry(root)
entry_cpf.grid(row=2, column=1, padx=10, pady=5)

label_endereco = tk.Label(root, text="Endereço Completo:")
label_endereco.grid(row=3, column=0, sticky="w", padx=10, pady=5)
entry_endereco = tk.Entry(root)
entry_endereco.grid(row=3, column=1, padx=10, pady=5)

label_cidade = tk.Label(root, text="Cidade de Nascimento:")
label_cidade.grid(row=4, column=0, sticky="w", padx=10, pady=5)
entry_cidade = tk.Entry(root)
entry_cidade.grid(row=4, column=1, padx=10, pady=5)

label_estado = tk.Label(root, text="Estado de Nascimento:")
label_estado.grid(row=5, column=0, sticky="w", padx=10, pady=5)
entry_estado = tk.Entry(root)
entry_estado.grid(row=5, column=1, padx=10, pady=5)

label_data = tk.Label(root, text="Data de Nascimento:")
label_data.grid(row=6, column=0, sticky="w", padx=10, pady=5)
entry_data = tk.Entry(root)
entry_data.grid(row=6, column=1, padx=10, pady=5)

label_pai = tk.Label(root, text="Nome do Pai:")
label_pai.grid(row=7, column=0, sticky="w", padx=10, pady=5)
entry_pai = tk.Entry(root)
entry_pai.grid(row=7, column=1, padx=10, pady=5)

label_mae = tk.Label(root, text="Nome da Mãe:")
label_mae.grid(row=8, column=0, sticky="w", padx=10, pady=5)
entry_mae = tk.Entry(root)
entry_mae.grid(row=8, column=1, padx=10, pady=5)

label_data_exp_rg = tk.Label(root, text="Data de Expedição do RG:")
label_data_exp_rg.grid(row=9, column=0, sticky="w", padx=10, pady=5)
entry_data_exp_rg = tk.Entry(root)
entry_data_exp_rg.grid(row=9, column=1, padx=10, pady=5)

# Botão para coletar os dados
btn_gerar = tk.Button(root, text="Gerar Dados", command=coletar_dados)
btn_gerar.grid(row=10, column=0, columnspan=2, pady=10)

# Iniciar a interface gráfica
root.mainloop()